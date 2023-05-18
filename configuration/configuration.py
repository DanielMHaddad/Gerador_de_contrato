from docx import Document
from docx2pdf import convert

# A classe Contrato encapsula a classe Document para criar um objeto Contrato que herda todos os atributos de Document.
#O encapsulamento foi necessário pois a classe "Document()" não perimite que ela sera herdada  
class Contrato:
    def __init__(self, caminho=None):
        self.documento = Document(caminho) if caminho else Document()

    # Função para ler um documento a partir de um caminho especificado.
    def read_doc(self, caminho):
        self.documento = Document(caminho)

    # Função para salvar o documento com um nome e caminho especificados.
    def save_doc(self, nome, ap, pasta_ap):
        self.documento.save(f'Contratos\{pasta_ap}\Contrato_{ap} - {nome}.docx')

    # Função para converter o documento em PDF com um nome e caminho especificados.
    def convert_to_pdf(self, nome, ap, pasta_ap):
        caminho_docx = f'Contratos\{pasta_ap}\Contrato_{ap} - {nome}.docx'
        caminho_pdf = f'Contratos\{pasta_ap}\Contrato_{ap} - {nome}.pdf'
        convert(caminho_docx, caminho_pdf)
        print(caminho_pdf)
